"""Word 文档文本提取"""
import logging
from pathlib import Path

logger = logging.getLogger(__name__)


def is_word_file(path: Path) -> bool:
    """判断是否为 Word 文档"""
    return path.suffix.lower() in [".docx", ".doc"]


def extract_text_from_docx(path: Path) -> str:
    """使用 python-docx 提取 Word 文档文本（段落 + 表格）"""
    from docx import Document

    doc = Document(str(path))
    texts = []

    # 提取段落
    for para in doc.paragraphs:
        if para.text.strip():
            texts.append(para.text)

    # 提取表格内容
    for table in doc.tables:
        for row in table.rows:
            row_text = " ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                texts.append(row_text)

    result = '\n'.join(texts)
    logger.info("extract_text_from_docx: file=%s chars=%d", path.name, len(result))
    return result


def extract_text_from_doc(path: Path) -> str:
    """提取 .doc 格式文本（需要转换为 .docx 或使用其他库）"""
    # .doc 格式处理较复杂，建议用户转换为 .docx
    # 这里返回提示信息
    msg = f"[.doc 格式暂不支持，请将文件转换为 .docx 格式: {path.name}]"
    logger.warning(msg)
    return msg
